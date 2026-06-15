from docx import Document
from openpyxl import load_workbook
import re
import copy

# ─── CẤU HÌNH ────────────────────────────────────────────────────────────────
EXCEL_PATH = r"d:\tool-hub\public\samples\Danh_Sach_Tham_Gia-Mau.xlsx"
DOCX_PATH  = r"d:\tool-hub\public\samples\PGE-Mau_Giay-Mau.docx"
# ─────────────────────────────────────────────────────────────────────────────


def get_start_date_from_excel(excel_path: str) -> str:
    """
    Đọc cột 'Thời gian bắt đầu' từ dòng đầu tiên có dữ liệu.
    Trả về chuỗi định dạng YYYYMMDD, ví dụ: '20260605'.

    Quy tắc đọc ngày (theo thứ tự ưu tiên):
    1. Nếu cell là TEXT (str)  → parse theo DD/MM/YYYY (anh tự nhập tay)
    2. Nếu cell là DATETIME    → đọc number_format để biết chiều DD/MM hay MM/DD
       - number_format bắt đầu bằng 'd' hoặc chứa 'dd/mm' → ngày trước tháng
       - Không xác định được → báo lỗi, yêu cầu nhập dạng text
    """
    from datetime import datetime

    # Cần load không read_only để đọc được number_format
    wb = load_workbook(excel_path, read_only=False, data_only=True)
    ws = wb.active

    # Tìm index cột từ header
    headers = [str(cell.value).strip() if cell.value else "" for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    try:
        col_idx = headers.index("Thời gian bắt đầu")
    except ValueError:
        raise ValueError(f"Không tìm thấy cột 'Thời gian bắt đầu'.\nHeaders hiện có: {headers}")

    for row in ws.iter_rows(min_row=2):
        cell = row[col_idx]
        val  = cell.value
        if val is None:
            continue

        # ── Trường hợp 1: Anh nhập text "05/06/2026" ──────────────────────────
        if isinstance(val, str):
            val = val.strip()
            try:
                dt = datetime.strptime(val, "%d/%m/%Y")
                return dt.strftime("%Y%m%d")
            except ValueError:
                pass
            # Thử thêm ISO format phòng khi ai đó nhập "2026-06-05"
            try:
                dt = datetime.strptime(val, "%Y-%m-%d")
                return dt.strftime("%Y%m%d")
            except ValueError:
                pass
            raise ValueError(
                f"Không parse được ngày từ text '{val}'.\n"
                f"Hãy nhập đúng định dạng DD/MM/YYYY (ví dụ: 05/06/2026)."
            )

        # ── Trường hợp 2: Cell là datetime (Excel tự chuyển) ──────────────────
        if hasattr(val, "strftime"):
            fmt = (cell.number_format or "").lower()
            # number_format dạng dd/mm/yyyy hoặc d/m/yyyy → ngày trước tháng
            if fmt.startswith("d") or "dd/mm" in fmt or "d/m" in fmt:
                # Giá trị trong datetime đã bị Excel hoán đổi tháng/ngày
                # → đảo lại: lấy month làm ngày, day làm tháng
                try:
                    corrected = datetime(val.year, val.day, val.month)
                    return corrected.strftime("%Y%m%d")
                except ValueError:
                    pass  # Ngày không hợp lệ sau khi đảo → rơi xuống lỗi

            raise ValueError(
                f"Ô '{cell.coordinate}' chứa ngày dạng datetime ({val.strftime('%d/%m/%Y')}) "
                f"nhưng không xác định được chiều ngày/tháng (number_format='{cell.number_format}').\n"
                f"→ Hãy đổi ô đó sang định dạng TEXT và nhập lại: 05/06/2026"
            )

    raise ValueError("Cột 'Thời gian bắt đầu' không có dữ liệu hợp lệ.")


def get_row_count_from_excel(excel_path: str) -> int:
    """Đếm số dòng dữ liệu (bỏ qua header)."""
    wb = load_workbook(excel_path, read_only=True)
    ws = wb.active
    count = 0
    first = True
    for row in ws.iter_rows(values_only=True):
        if first:
            first = False
            continue
        # Dòng có ít nhất 1 ô không rỗng
        if any(c is not None for c in row):
            count += 1
    return count


def replace_text_in_run(run, pattern: str, replacement: str):
    """Thay thế text trong một run, giữ nguyên định dạng."""
    if pattern in run.text:
        run.text = run.text.replace(pattern, replacement)
        return True
    return False


def replace_in_paragraph(para, pattern: str, replacement: str) -> bool:
    """
    Thay thế pattern trong paragraph, cố gắng giữ formatting của runs.
    Nếu pattern nằm trải qua nhiều run, dùng fallback gộp text.
    """
    # Thử thay trong từng run trước
    if any(pattern in run.text for run in para.runs):
        changed = False
        for run in para.runs:
            if replace_text_in_run(run, pattern, replacement):
                changed = True
        return changed

    # Pattern có thể bị tách qua nhiều run → gộp lại kiểm tra
    full_text = para.text
    if pattern not in full_text:
        return False

    # Xây lại text của run đầu tiên, xóa các run còn lại
    new_text = full_text.replace(pattern, replacement)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    return True


def process_docx(docx_path: str, date_str: str, row_count: int):
    """
    Duyệt toàn bộ docx, thay thế:
    - {{So_GCN}}  →  PGE-{date_str}-{TT}   (TT lấy từ vị trí trang/phần)
    - {{TT}}      →  số thứ tự tương ứng

    Cấu trúc file mẫu: mỗi section/bảng tương ứng 1 học viên.
    Script sẽ:
    1. Đếm số lần xuất hiện của {{So_GCN}} hoặc pattern PGE-YYYYMMDD-{{TT}}
    2. Thay thế lần lượt theo TT từ 1 → row_count
    """
    doc = Document(docx_path)

    # Pattern cần thay
    PATTERN_SERIAL = re.compile(r'PGE-\d{8}-\d+')   # ví dụ PGE-20260506-1
    PLACEHOLDER    = "{{So_GCN}}"
    PLACEHOLDER_TT = "{{TT}}"

    counter = [0]  # dùng list để closure có thể thay đổi

    def make_so_gcn(tt: int) -> str:
        return f"PGE-{date_str}-{tt}"

    def process_para(para) -> bool:
        text = para.text
        changed = False

        # Trường hợp 1: placeholder {{So_GCN}}
        if PLACEHOLDER in text:
            counter[0] += 1
            changed = replace_in_paragraph(para, PLACEHOLDER, make_so_gcn(counter[0]))

        # Trường hợp 2: pattern PGE-YYYYMMDD-N cũ (cập nhật ngày + giữ số TT)
        elif PATTERN_SERIAL.search(text):
            def replacer(m):
                # Giữ nguyên số TT cuối
                old = m.group()
                tt = int(old.rsplit("-", 1)[-1])
                return make_so_gcn(tt)
            new_text = PATTERN_SERIAL.sub(replacer, text)
            if new_text != text:
                replace_in_paragraph(para, text, new_text)
                changed = True

        # Trường hợp 3: {{TT}} độc lập (ít gặp, nhưng xử lý luôn)
        if PLACEHOLDER_TT in para.text:
            # Tìm số TT gần nhất đã dùng
            tt = counter[0] if counter[0] > 0 else 1
            replace_in_paragraph(para, PLACEHOLDER_TT, str(tt))
            changed = True

        return changed

    # Duyệt paragraphs tự do
    for para in doc.paragraphs:
        process_para(para)

    # Duyệt tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)

    doc.save(docx_path)
    print(f"✅ Đã lưu file: {docx_path}")
    print(f"   Số GCN đã thay: {counter[0]}")
    print(f"   Số học viên trong Excel: {row_count}")
    if counter[0] != row_count:
        print(f"   ⚠️  Số lượng không khớp — kiểm tra lại file mẫu.")


if __name__ == "__main__":
    try:
        print("📖 Đọc ngày bắt đầu từ Excel...")
        date_str = get_start_date_from_excel(EXCEL_PATH)
        print(f"   Thời gian bắt đầu: {date_str[:4]}-{date_str[4:6]}-{date_str[6:]} → chuỗi GCN: {date_str}")

        print("📊 Đếm số học viên...")
        row_count = get_row_count_from_excel(EXCEL_PATH)
        print(f"   Số dòng dữ liệu: {row_count}")

        print("✏️  Đang cập nhật file docx...")
        process_docx(DOCX_PATH, date_str, row_count)

    except Exception as e:
        print(f"❌ Lỗi: {e}")
        raise
