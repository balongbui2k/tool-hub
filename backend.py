"""
PGE Backend Server
==================
FastAPI backend cho ToolHub - PGE Certificate Suite.
Port lại logic từ tool.py để chạy qua HTTP API.

Cài đặt:
    pip install fastapi uvicorn python-multipart pandas openpyxl docxtpl docx2pdf docxcompose PyPDF2 python-docx

Chạy:
    python backend.py
    hoặc: uvicorn backend:app --reload --port 8000
"""

import io
import os
import re
import sys
import copy
import json
import zipfile
import tempfile
import unicodedata
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

import pandas as pd
import PyPDF2
import openpyxl
from openpyxl.styles import Alignment, Border, Side, Font, PatternFill
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
try:
    from docx2pdf import convert as _docx2pdf_convert
except ImportError:
    _docx2pdf_convert = None
from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
import socket
import uuid
import shutil

app = FastAPI(title="PGE Tool Backend", version="2.0.0")

@app.get("/")
async def root():
    return {
        "message": "Phúc Gia Hub API is running!",
        "status": "online",
        "docs": "/docs"
    }

# Allow Public Origins for Production
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # In production, you might want to restrict this to your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition", "x-file-count", "x-logs"],
)

# Thư mục lưu file upload
UPLOAD_DIR = Path("public/files")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
app.mount("/files", StaticFiles(directory="public/files"), name="files")

def get_local_ip():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        # Connect to a public DNS to determine default route
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except:
        return "127.0.0.1"

@app.post("/api/upload")
async def upload_file(request: Request, file: UploadFile = File(...)):
    try:
        # Prevent tracking issues, store file with a safe unique name
        ext = Path(file.filename).suffix
        safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename.replace(' ', '_')}"
        file_path = UPLOAD_DIR / safe_name
        
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        # DYNAMIC BASE URL: detect how the user is accessing the server
        # This ensures QR codes work for both localhost and Network IP
        base_url = str(request.base_url).rstrip('/')
        file_url = f"{base_url}/api/download/{safe_name}"
        
        print(f"📁 File uploaded: {file.filename} -> {file_url}")
        
        return {"status": "success", "url": file_url, "filename": file.filename}
    except Exception as e:
        print(f"❌ Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    # Trích xuất tên file gốc (Bỏ đoạn uuid ngẫu nhiên ở phần đầu để tên file tải về đẹp hơn)
    download_name = filename
    if "_" in filename:
        parts = filename.split("_", 1)
        if len(parts[0]) == 8:
            download_name = parts[1]
            
    return FileResponse(
        path=file_path, 
        filename=download_name, 
        headers={"Content-Disposition": f'attachment; filename="{download_name}"'}
    )

# ── Utils (ported from tool.py) ────────────────────────────────────────────

def viet_to_ascii(text: str) -> str:
    text = str(text).replace('Đ', 'D').replace('đ', 'd')
    text = unicodedata.normalize('NFD', text)
    return ''.join(c for c in text if unicodedata.category(c) != 'Mn')

def name_to_slug(full_name: str) -> str:
    ascii_name = viet_to_ascii(str(full_name).strip())
    return re.sub(r'[^A-Za-z0-9]+', '_', ascii_name).strip('_')

def course_to_slug(course: str) -> str:
    ascii_course = viet_to_ascii(str(course).strip().title())
    return re.sub(r'[^A-Za-z0-9]+', '_', ascii_course).strip('_')

def format_date_str(val):
    if pd.isna(val) or val is None: return ""
    if isinstance(val, (datetime, pd.Timestamp)): 
        return val.strftime("%d/%m/%Y")
    
    s = str(val).strip()
    if not s or s.lower() == 'nan': return ""
    if " " in s: s = s.split(" ")[0]
    
    # Ưu tiên parse Ngày/Tháng/Năm (VN) bằng pandas
    try:
        dt = pd.to_datetime(s, dayfirst=True, errors='coerce')
        if pd.notna(dt):
            return dt.strftime("%d/%m/%Y")
    except:
        pass
        
    return s

def format_date_range(start_val, end_val):
    s = format_date_str(start_val)
    e = format_date_str(end_val)
    if not s: return ""
    if not e or s == e: return s
    return f"{s} - {e}"

def convert_docx_to_pdf_win32(in_path: str, out_path: str) -> bool:
    """
    Chuyển đổi DOCX → PDF bằng Win32COM (comtypes), đáng tin cậy nhất trên Windows.
    Trả về True nếu thành công, False nếu thất bại.
    """
    try:
        import comtypes.client
        # wdFormatPDF = 17
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(in_path)
        doc.SaveAs(out_path, FileFormat=17)
        doc.Close()
        word.Quit()
        return os.path.exists(out_path)
    except Exception as e:
        print(f"❌ convert_docx_to_pdf_win32 lỗi: {e}")
        return False

def convert_docx_to_pdf_linux(in_path: str, out_dir: str) -> bool:
    """Chuyển đổi DOCX → PDF bằng LibreOffice (soffice), chuẩn cho Linux/Docker."""
    try:
        import subprocess
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', out_dir, in_path],
            check=True,
            capture_output=True
        )
        return True
    except Exception as e:
        print(f"❌ convert_docx_to_pdf_linux lỗi: {e}")
        return False

def extract_first_date_for_prefix(text):
    match = re.search(r'\d{4}[./-]\d{2}[./-]\d{2}', text)
    if match: return match.group().replace("-", ".").replace("/", ".")
    match = re.search(r'\d{2}[./-]\d{2}[./-]\d{4}', text)
    if match:
        d = match.group().replace("-", "/").replace(".", "/")
        try: return datetime.strptime(d, "%d/%m/%Y").strftime("%Y.%m.%d")
        except: return text
    return text

def sanitize_docx(doc):
    """Xóa cài đặt Mail Merge khỏi file Word để tránh lỗi pop-up khi mở file"""
    try:
        settings = doc.settings.element
        removed = False
        # Danh sách các thẻ liên quan đến Mail Merge cần xóa
        for child in list(settings):
            if 'mailMerge' in child.tag:
                settings.remove(child)
                removed = True
        if removed:
            print("🧹 Đã xóa cài đặt Mail Merge khỏi file Word mẫu.")
        return removed
    except Exception:
        pass
    return False

# Mapping các tên cột có thể có trong Excel
COL_MAP = {
    'name': ['họ và tên', 'họ tên', 'full name', 'tên học viên', 'tên'],
    'course': ['khóa học', 'tên khóa học', 'khóa', 'course name'],
    'start_date': ['thời gian bắt đầu', 'ngày bắt đầu', 'từ ngày', 'start date', 'tg bắt đầu'],
    'end_date': ['thời gian kết thúc', 'ngày kết thúc', 'đến ngày', 'end date', 'tg kết thúc'],
    'gender': ['giới tính', 'gender', 'phái'],
    'birth_date': ['ngày sinh', 'năm sinh', 'birth date', 'ngay sinh'],
    'sign_date': ['ngày ký', 'ngày ký tên', 'sign date', 'ngày cấp'],
    'title': ['danh xưng', 'ông/bà', 'title', 'chức danh'],
}

def get_col_val(row, key, default=None):
    """Lấy giá trị từ row dựa trên mapping cột"""
    cols = [str(c).strip().lower() for c in row.index]
    for candidate in COL_MAP.get(key, []):
        if candidate in cols:
            return row[candidate]
    return default

def read_excel_names(excel_bytes: bytes):
    """Đọc Excel từ bytes, ưu tiên đọc sheet đang active giống như các bước trước để đồng bộ dữ liệu"""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(excel_bytes), data_only=True, read_only=True)
    target_sh = wb.active.title
    wb.close()
    
    with pd.ExcelFile(io.BytesIO(excel_bytes), engine='openpyxl') as xl:
        if target_sh in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=target_sh)
        else:
            df = pd.read_excel(xl, sheet_name=xl.sheet_names[0])
            
    df.columns = [str(c).strip().lower() for c in df.columns]

    # Tìm thông tin prefix từ dòng đầu tiên
    first_row = df.iloc[0]
    first_row_start = get_col_val(first_row, 'start_date')
    
    if isinstance(first_row_start, (datetime, pd.Timestamp)):
        s_date_prefix = first_row_start.strftime("%Y.%m.%d")
    else:
        s_date_prefix = extract_first_date_for_prefix(str(first_row_start))

    c_name_default = str(get_col_val(first_row, 'course', '')).strip()
    prefix = f"{s_date_prefix}-{course_to_slug(c_name_default)}"

    names_original, names_slug, rows_data = [], [], []
    for _, row in df.iterrows():
        name = get_col_val(row, 'name')
        if not name or str(name).lower() == 'nan': continue
        names_original.append(str(name))
        names_slug.append(name_to_slug(str(name)))
        rows_data.append(row)

    return prefix, names_slug, names_original, rows_data, c_name_default, s_date_prefix


# ── API Endpoints ──────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "service": "PGE Tool Backend v2.0"}


@app.post("/api/cert-suite")
async def cert_suite(
    excel: UploadFile = File(...),
    docx: UploadFile = File(...),
    output_name: str = Form("Tong_Hop_Chung_Chi.docx"),
    output_format: str = Form("docx"),
):
    """Tool 1: Tạo file Word tổng hợp chứng chỉ"""
    logs = []

    def log(type_, text):
        logs.append({"type": type_, "text": text})

    try:
        excel_bytes = await excel.read()
        docx_bytes = await docx.read()

        # Kiểm tra nhanh định dạng file để tránh lỗi ValueError khi chọn nhầm file Excel vào ô Word
        if not docx.filename.lower().endswith(".docx"):
            return JSONResponse({"logs": logs, "error": f"File mẫu '{docx.filename}' không phải là định dạng Word (.docx). Vui lòng kiểm tra lại!"}, status_code=400)
        
        # Thử kiểm tra nội dung file bằng cách đọc header (optional nhưng an toàn)
        if b"word/document.xml" not in docx_bytes[:8000] and b"word/_rels" not in docx_bytes[:8000]:
             # Nếu file chứa 'spreadsheetml', chắc chắn là chọn nhầm file Excel
             if b"spreadsheetml" in docx_bytes[:8000]:
                 return JSONResponse({"logs": logs, "error": "Bạn đã chọn nhầm file Excel vào mục 'Biểu mẫu Word'. Vui lòng chọn lại đúng file .docx!"}, status_code=400)

        log("info", "📚 Đang đọc dữ liệu từ Excel...")
        prefix, names_slug, names_original, rows_data, c_name_default, s_date_prefix = read_excel_names(excel_bytes)
        log("info", f"✓ Tìm thấy {len(names_original)} học viên. Prefix: {prefix}")

        # Override output_name theo yêu cầu
        safe_course = _td_make_filename_safe(c_name_default) if c_name_default else "Khong_Ten"
        output_name = f"{s_date_prefix}-Chung_Chi-{safe_course}.docx"

        doc_objs = []

        for idx, row in enumerate(rows_data):
            name = names_original[idx]
            c_name = str(get_col_val(row, 'course', c_name_default)).strip()
            if not c_name or c_name.lower() == 'nan': c_name = c_name_default

            dx = str(get_col_val(row, 'title', '')).strip()
            if not dx or dx.lower() == 'nan':
                gt = str(get_col_val(row, 'gender', '')).strip().lower()
                if gt in ['nam', 'm', 'male']: dx = 'Ông'
                elif gt in ['nữ', 'nu', 'f', 'female']: dx = 'Bà'
                else: dx = 'Ông/Bà'

            ns = format_date_str(get_col_val(row, 'birth_date', ''))

            t_starts = get_col_val(row, 'start_date')
            t_ends = get_col_val(row, 'end_date')
            tg_hoc = format_date_range(t_starts, t_ends)

            raw_sign = get_col_val(row, 'sign_date', '')
            sign_t = ""
            sign_d = None
            if isinstance(raw_sign, datetime):
                sign_d = raw_sign
                sign_t = f"Hà Nội, ngày {raw_sign.day:02d} tháng {raw_sign.month:02d} năm {raw_sign.year}"
            elif pd.notna(raw_sign) and str(raw_sign).strip() and str(raw_sign).lower() != 'nan':
                match = re.search(r'(\d{1,2})\D+(\d{1,2})\D+(\d{4})', str(raw_sign))
                if match:
                    d, m, y = int(match.group(1)), int(match.group(2)), int(match.group(3))
                    sign_d = datetime(y, m, d)
                    sign_t = f"Hà Nội, ngày {d:02d} tháng {m:02d} năm {y}"

            if not sign_t or not sign_d:
                ref = pd.to_datetime(t_starts, errors='coerce', dayfirst=True)
                if pd.isna(ref):
                    try: ref = datetime.strptime(s_date_prefix, "%Y.%m.%d")
                    except: ref = datetime.now()
                sign_d = ref + timedelta(days=7)
                sign_t = f"Hà Nội, ngày {sign_d.day:02d} tháng {sign_d.month:02d} năm {sign_d.year}"

            # So_GCN logic: PGE-YYYYMMDD-TT
            so_gcn = f"PGE-{sign_d.strftime('%Y%m%d')}-{idx+1:02d}"

            tpl = DocxTemplate(io.BytesIO(docx_bytes))
            
            # Xóa cài đặt Mail Merge khỏi template trước khi render
            sanitize_docx(tpl.docx)
            
            tpl.render({
                'Danh_xưng': str(dx),
                'Họ_và_tên': str(name),
                'Ngày_sinh': str(ns),
                'Khóa_học': c_name,
                'Thời_gian': tg_hoc,
                'TT': f"{idx+1:02d}",
                'So_GCN': so_gcn,
                'Ngay_Ki': sign_t,
            })

            mem = io.BytesIO()
            tpl.save(mem)
            mem.seek(0)
            doc_objs.append(Document(mem))
            log("info", f"  ✓ {name} | {tg_hoc}")

        if not doc_objs:
            return JSONResponse({"logs": logs, "error": "Không tìm thấy học viên nào trong Excel!"}, status_code=400)

        log("info", f"\n📎 Đang ghép {len(doc_objs)} chứng chỉ thành 1 file...")
        master = doc_objs[0]
        comp = Composer(master)
        for i in range(1, len(doc_objs)):
            master.add_page_break()
            comp.append(doc_objs[i])

        sanitize_docx(comp.doc)

        out_buf = io.BytesIO()
        comp.save(out_buf)
        out_buf.seek(0)

        log("success", f"✅ Xong! Đã tạo file tổng hợp với {len(doc_objs)} chứng chỉ.")

        file_bytes = out_buf.getvalue()
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        if output_format.lower() == "pdf":
            log("info", "📄 Đang chuyển đổi sang PDF bằng Microsoft Word COM...")
            try:
                temp_id = uuid.uuid4().hex[:8]
                in_path = str((UPLOAD_DIR / f"temp_{temp_id}.docx").absolute())
                out_path = str((UPLOAD_DIR / f"temp_{temp_id}.pdf").absolute())

                with open(in_path, "wb") as f:
                    f.write(file_bytes)
                log("info", f"  ↳ File tạm DOCX: {in_path}")

                # ── XỬ LÝ CHUYỂN ĐỔI PDF ──
                success = False
                if sys.platform == "win32":
                    # Phương án 1: Win32COM (comtypes) - đáng tin cậy nhất trên Windows
                    success = convert_docx_to_pdf_win32(in_path, out_path)

                    # Phương án 2: Fallback docx2pdf nếu Win32COM thất bại
                    if not success and _docx2pdf_convert is not None:
                        log("warn", "  ⚠ Win32COM thất bại, thử docx2pdf...")
                        try:
                            _docx2pdf_convert(in_path, out_path)
                            success = os.path.exists(out_path)
                        except Exception as e2:
                            log("error", f"  ❌ docx2pdf cũng thất bại: {e2}")
                else:
                    # Phương án cho Linux/Docker: LibreOffice (soffice)
                    log("info", "  📄 Đang chuyển đổi bằng LibreOffice (soffice)...")
                    success = convert_docx_to_pdf_linux(in_path, str(UPLOAD_DIR.absolute()))
                    
                    # LibreOffice mặc định tạo file tên giống file gốc nhưng đuôi .pdf
                    expected_out = in_path.replace(".docx", ".pdf")
                    if os.path.exists(expected_out):
                        # Đổi tên về out_path (temp_xxxx.pdf)
                        if os.path.exists(out_path): os.remove(out_path)
                        os.rename(expected_out, out_path)
                        success = True

                if success:
                    with open(out_path, "rb") as f:
                        file_bytes = f.read()
                    media_type = "application/pdf"
                    if output_name.lower().endswith(".docx"):
                        output_name = output_name[:-5] + ".pdf"
                    elif not output_name.lower().endswith(".pdf"):
                        output_name += ".pdf"
                    log("success", f"✅ Chuyển đổi PDF thành công! ({len(file_bytes)//1024} KB)")
                else:
                    log("error", "❌ Không thể tạo PDF. Trên Windows hãy kiểm tra MS Word, trên Linux hãy kiểm tra LibreOffice.")

                # Dọn file tạm
                for p in [in_path, out_path]:
                    if os.path.exists(p):
                        try: os.remove(p)
                        except: pass

            except Exception as e:
                import traceback
                log("error", f"❌ Lỗi chuyển đổi PDF: {str(e)}")
                log("error", traceback.format_exc())

        return Response(
            content=file_bytes,
            media_type=media_type,
            headers={
                "Content-Disposition": f'attachment; filename="{output_name}"',
                "x-logs": str(len(logs)),
            }
        )

    except Exception as e:
        import traceback
        log("error", f"❌ Lỗi: {str(e)}")
        log("error", traceback.format_exc())
        return JSONResponse({"logs": logs, "error": str(e)}, status_code=500)


@app.post("/api/split-pdf")
async def split_pdf(
    excel: UploadFile = File(...),
    pdf: UploadFile = File(...),
    drive_link: Optional[str] = Form(None),
):
    """Tool 2: Tách PDF theo danh sách học viên"""
    logs = []

    def log(type_, text):
        logs.append({"type": type_, "text": text})

    try:
        excel_bytes = await excel.read()
        pdf_bytes = await pdf.read()

        log("info", "📚 Đang đọc danh sách từ Excel...")
        prefix, names_slug, names_original, _, _, _ = read_excel_names(excel_bytes)
        log("info", f"✓ {len(names_original)} học viên. Prefix: {prefix}")

        work_pdf_bytes = pdf_bytes

        # If uploaded file is DOCX, we'd need to convert - skip for now, require PDF
        if pdf.filename and pdf.filename.lower().endswith('.docx'):
            log("warn", "⚠ File DOCX trực tiếp không được hỗ trợ trong web mode. Vui lòng chuyển sang PDF trước.")
            return JSONResponse({"logs": logs, "error": "Vui lòng upload file PDF (không phải DOCX). Hãy chuyển Word→PDF bằng tool Tool 1 trước."}, status_code=400)

        log("info", "✂️  Đang tách từng trang PDF...")

        zip_buf = io.BytesIO()
        count = 0

        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            reader = PyPDF2.PdfReader(io.BytesIO(work_pdf_bytes))
            total_pages = len(reader.pages)

            for idx in range(min(total_pages, len(names_slug))):
                writer = PyPDF2.PdfWriter()
                writer.add_page(reader.pages[idx])

                new_name = f"{prefix}-{names_slug[idx]}.pdf"
                page_buf = io.BytesIO()
                writer.write(page_buf)
                zf.writestr(new_name, page_buf.getvalue())

                log("info", f"  ➜ {new_name}")
                count += 1

        if drive_link:
            log("info", f"🔗 Google Drive: {drive_link}")

        log("success", f"✅ Xong! Đã tách {count} file PDF thành công.")

        zip_buf.seek(0)
        return Response(
            content=zip_buf.getvalue(),
            media_type="application/zip",
            headers={
                "Content-Disposition": 'attachment; filename="Chung_Chi_PDF.zip"',
                "x-file-count": str(count),
            }
        )

    except Exception as e:
        import traceback
        log("error", f"❌ Lỗi: {str(e)}")
        return JSONResponse({"logs": logs, "error": str(e)}, status_code=500)


@app.post("/api/rename")
async def rename_pdfs(
    excel: UploadFile = File(...),
    zip: UploadFile = File(...),
    drive_link: Optional[str] = Form(None),
):
    """Tool 3: Đổi tên file PDF theo danh sách"""
    logs = []

    def log(type_, text):
        logs.append({"type": type_, "text": text})

    try:
        excel_bytes = await excel.read()
        zip_bytes = await zip.read()

        log("info", "📚 Đang đọc danh sách từ Excel...")
        prefix, names_slug, _, _, _, _ = read_excel_names(excel_bytes)
        log("info", f"✓ {len(names_slug)} học viên. Prefix: {prefix}")

        log("info", "🏷️ Đang đổi tên file trong ZIP...")

        out_zip_buf = io.BytesIO()
        count = 0

        with zipfile.ZipFile(io.BytesIO(zip_bytes), 'r') as zin:
            with zipfile.ZipFile(out_zip_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
                pdf_files = sorted(
                    [f for f in zin.namelist() if f.lower().endswith('.pdf')],
                    key=lambda x: (
                        int(re.search(r'-(\d+)\.pdf$', x).group(1))
                        if re.search(r'-(\d+)\.pdf$', x)
                        else 9999
                    )
                )

                for item in zin.namelist():
                    if not item.lower().endswith('.pdf'):
                        continue
                    name_base = item.rsplit('.pdf', 1)[0]
                    parts = name_base.split('-')
                    idx_str = parts[-1]
                    if idx_str.isdigit():
                        idx = int(idx_str)
                        if 1 <= idx <= len(names_slug):
                            new_name = f"{prefix}-{names_slug[idx-1]}.pdf"
                            zout.writestr(new_name, zin.read(item))
                            log("info", f"  ➜ {new_name}")
                            count += 1
                            continue
                    # Keep original if no match
                    zout.writestr(item, zin.read(item))

        if drive_link:
            log("info", f"🔗 Google Drive: {drive_link}")

        log("success", f"✅ Xong! Đã đổi tên {count} file PDF.")

        out_zip_buf.seek(0)
        return Response(
            content=out_zip_buf.getvalue(),
            media_type="application/zip",
            headers={
                "Content-Disposition": 'attachment; filename="Chung_Chi_DaDoiTen.zip"',
                "x-file-count": str(count),
            }
        )

    except Exception as e:
        import traceback
        log("error", f"❌ Lỗi: {str(e)}")
        return JSONResponse({"logs": logs, "error": str(e)}, status_code=500)


# ── Training Dossier (Tạo Hồ Sơ Đào Tạo) ──────────────────────────────────

def _td_normalize_text(text):
    if text is None: return ""
    return unicodedata.normalize('NFC', str(text).strip())

def _td_clean_for_match(text):
    t = _td_normalize_text(text)
    return re.sub(r'\s+', '', t).lower()

def _td_parse_date(val):
    """Parse date from Excel cell value, swapping day/month for VN format."""
    if val is None: return None
    if isinstance(val, datetime):
        try:
            return datetime(val.year, val.day, val.month, val.hour, val.minute, val.second)
        except ValueError:
            return val
    if isinstance(val, (int, float)):
        dt = datetime(1899, 12, 30) + timedelta(days=val)
        try:
            return datetime(dt.year, dt.day, dt.month, dt.hour, dt.minute, dt.second)
        except ValueError:
            return dt
    if isinstance(val, str):
        val = val.strip()
        for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%Y-%m-%d"):
            try: return datetime.strptime(val, fmt)
            except ValueError: continue
    return None

def _td_make_filename_safe(name):
    ascii_name = viet_to_ascii(str(name).strip())
    ascii_name = ascii_name.title()
    ascii_name = re.sub(r'[^A-Za-z0-9]', '_', ascii_name)
    ascii_name = re.sub(r'_+', '_', ascii_name).strip('_')
    return ascii_name[:80]

def _td_make_rPr(bold=False, italic=False):
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), 'Times New Roman')
    fonts.set(qn('w:hAnsi'), 'Times New Roman')
    fonts.set(qn('w:cs'), 'Times New Roman')
    rPr.append(fonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    if italic:
        rPr.append(OxmlElement('w:i'))
        rPr.append(OxmlElement('w:iCs'))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '26')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '26')
    rPr.append(szCs)
    return rPr

def _td_set_cell_mixed(cell, segments):
    """segments: list of (text, bold, italic) tuples"""
    para = cell.paragraphs[0]
    for run in para.runs: run._r.getparent().remove(run._r)
    for seg in segments:
        text = str(seg[0])
        bold = seg[1]
        italic = seg[2] if len(seg) > 2 else False
        r = OxmlElement('w:r')
        r.append(_td_make_rPr(bold=bold, italic=italic))
        lines = text.split('\n')
        for idx, line in enumerate(lines):
            if idx > 0:
                br = OxmlElement('w:br')
                r.append(br)
            t = OxmlElement('w:t')
            t.text = line
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
        para._p.append(r)

def _td_parse_excel(excel_bytes: bytes):
    """Parse the master Excel file and return header info + column mapping."""
    wb = openpyxl.load_workbook(io.BytesIO(excel_bytes), data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    header = [str(h).strip() if h else "" for h in rows[0]]
    col_map = {h: i for i, h in enumerate(header)}
    return wb, ws, rows, header, col_map


@app.post("/api/training-dossier/courses")
async def td_get_courses(excel: UploadFile = File(...)):
    """Quét file Excel Danh Sách Tham Gia và trả về danh sách khóa học."""
    try:
        excel_bytes = await excel.read()
        _, _, rows, header, col_map = _td_parse_excel(excel_bytes)

        IDX_TEN_LOP = col_map.get("Khóa học", col_map.get("Tên lớp học", -1))
        IDX_NGAY_BD = col_map.get("Thời gian bắt đầu", col_map.get("Ngày bắt đầu", -1))
        IDX_NGAY_KT = col_map.get("Thời gian kết thúc", col_map.get("Ngày kết thúc", -1))
        IDX_GIANG_VIEN = col_map.get("Giảng viên", -1)

        if IDX_TEN_LOP == -1 or IDX_NGAY_BD == -1:
            raise HTTPException(status_code=400, detail="Không tìm thấy cột 'Khóa học' hoặc 'Thời gian bắt đầu' trong file Excel.")

        courses = {}
        for row in rows[1:]:
            if len(row) <= max(IDX_TEN_LOP, IDX_NGAY_BD): continue

            ten_lop = _td_normalize_text(row[IDX_TEN_LOP]) if row[IDX_TEN_LOP] else "N/A"
            ngay_bd = _td_parse_date(row[IDX_NGAY_BD])

            if not ngay_bd or ten_lop == "N/A": continue

            ngay_kt = _td_parse_date(row[IDX_NGAY_KT]) if IDX_NGAY_KT != -1 and IDX_NGAY_KT < len(row) else None
            giang_vien = ""
            if IDX_GIANG_VIEN != -1 and IDX_GIANG_VIEN < len(row):
                giang_vien = _td_normalize_text(row[IDX_GIANG_VIEN])

            key = (_td_clean_for_match(ten_lop), ngay_bd.strftime("%Y-%m-%d"))
            if key not in courses:
                courses[key] = {
                    "ten_lop": ten_lop,
                    "ngay_bd": ngay_bd.strftime("%Y-%m-%d"),
                    "ngay_kt": ngay_kt.strftime("%Y-%m-%d") if ngay_kt else "",
                    "giang_vien": giang_vien,
                    "count": 0
                }

            if not courses[key]["giang_vien"] and giang_vien:
                courses[key]["giang_vien"] = giang_vien

            courses[key]["count"] += 1

        for c in courses.values():
            if not c["giang_vien"]: c["giang_vien"] = "Chưa xác định"

        result = sorted(list(courses.values()), key=lambda x: x["ngay_bd"], reverse=True)
        return JSONResponse(content=result)

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


def _td_classify_unit(bo_phan: str) -> str:
    """Phân loại đơn vị: PGL, PGU, hoặc tên đơn vị ngoài."""
    bp = str(bo_phan).strip().upper()
    if "PGL" in bp:
        return "PGL"
    if "PGU" in bp:
        return "PGU"
    # Đơn vị ngoài - giữ nguyên tên gốc
    return str(bo_phan).strip() if bo_phan and str(bo_phan).strip() else "Khac"


def _td_generate_docx(tpl_docx_bytes: bytes, course: dict, participants: list,
                       dt_bd, dt_kt) -> io.BytesIO:
    """Tạo file BM03 DOCX cho một nhóm participants."""
    doc = Document(io.BytesIO(tpl_docx_bytes))

    # Fill table 0 (Header)
    t0 = doc.tables[0]
    date_range = f"{dt_bd.day:02d}/{dt_bd.month:02d}/{dt_bd.year}"
    if dt_kt and dt_kt.date() != dt_bd.date():
        date_range += " - " + f"{dt_kt.day:02d}/{dt_kt.month:02d}/{dt_kt.year}"
    _td_set_cell_mixed(t0.cell(0, 0), [("Ngày: ", True), (date_range, False)])

    gv = course.get("giang_vien", "")
    if gv and gv != "Chưa xác định":
        _td_set_cell_mixed(t0.cell(3, 1), [("Giảng viên: ", True), (gv, False)])

    _td_set_cell_mixed(t0.cell(1, 1), [(course["ten_lop"], False)])

    # Table 2 (Signatures)
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        hanoi_date = f"Hà Nội, ngày {dt_bd.day} tháng {dt_bd.month} năm {dt_bd.year}"
        _td_set_cell_mixed(t2.cell(0, 1), [(hanoi_date, False, True)])

        if gv and gv != "Chưa xác định":
            _td_set_cell_mixed(t2.cell(1, 1), [("GIẢNG VIÊN", True)])
            if len(t2.rows) > 3:
                _td_set_cell_mixed(t2.cell(3, 1), [(gv, False)])

    # Fill table 1 (Participants)
    t1 = doc.tables[1]
    template_row = t1.rows[1]
    template_tr = template_row._tr
    for i, p in enumerate(participants):
        new_tr = copy.deepcopy(template_tr)
        cells = new_tr.findall(qn('w:tc'))
        data_to_fill = [f"{i+1}.", p["ho_ten"], p["bo_phan"]]
        for idx, txt in enumerate(data_to_fill):
            if txt is not None and idx < len(cells):
                tc_para = cells[idx].findall(qn('w:p'))[0]

                if idx == 0:
                    pPrs = tc_para.findall(qn('w:pPr'))
                    if not pPrs:
                        pPr = OxmlElement('w:pPr')
                        tc_para.insert(0, pPr)
                    else:
                        pPr = pPrs[0]
                    for numPr in pPr.findall(qn('w:numPr')): pPr.remove(numPr)
                    for ind in pPr.findall(qn('w:ind')): pPr.remove(ind)
                    jcs = pPr.findall(qn('w:jc'))
                    if not jcs:
                        jc = OxmlElement('w:jc')
                        pPr.append(jc)
                    else:
                        jc = jcs[0]
                    jc.set(qn('w:val'), 'center')

                for r in tc_para.findall(qn('w:r')): tc_para.remove(r)
                nr = OxmlElement('w:r')
                nr.append(_td_make_rPr(bold=False))
                nt = OxmlElement('w:t')
                nt.text = str(txt)
                nr.append(nt)
                tc_para.append(nr)

        template_tr.getparent().insert(list(template_tr.getparent()).index(template_tr), new_tr)
    template_tr.getparent().remove(template_tr)

    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_buf.seek(0)
    return docx_buf


def _td_generate_xlsx(tpl_xlsx_bytes: bytes, course: dict, participants: list) -> io.BytesIO:
    """Tạo file XLSX danh sách điểm danh cho một nhóm participants."""
    wb_tpl = openpyxl.load_workbook(io.BytesIO(tpl_xlsx_bytes))
    ws_tpl = wb_tpl.active
    ws_tpl["F4"].value = course["ten_lop"]
    ws_tpl["F5"].value = "Ông Bùi Bá Long                       Điện thoại: 0327 706 996"

    # Capture row 7 styles
    row7_styles = []
    for col_idx in range(1, 11):
        cell = ws_tpl.cell(row=7, column=col_idx)
        row7_styles.append({
            "font": copy.copy(cell.font) if cell.font else None,
            "fill": copy.copy(cell.fill) if cell.fill else None,
            "border": copy.copy(cell.border) if cell.border else None,
            "alignment": copy.copy(cell.alignment) if cell.alignment else None,
            "number_format": cell.number_format,
        })
    row7_height = ws_tpl.row_dimensions[7].height if 7 in ws_tpl.row_dimensions else 25

    # Capture Lưu ý section
    luuy_header_font = copy.copy(ws_tpl["A10"].font)
    luuy_header_align = copy.copy(ws_tpl["A10"].alignment)
    luuy_notes_font = copy.copy(ws_tpl["A11"].font)
    luuy_notes_align = copy.copy(ws_tpl["A11"].alignment)
    luuy_header_height = ws_tpl.row_dimensions[10].height if 10 in ws_tpl.row_dimensions else 31.5
    luuy_notes_height = ws_tpl.row_dimensions[11].height if 11 in ws_tpl.row_dimensions else 43.5
    notes_text = ws_tpl["A11"].value

    # Unmerge template data rows
    ranges_to_remove = [
        str(r) for r in ws_tpl.merged_cells.ranges
        if r.min_row >= 7 and r.max_row <= 11
    ]
    for r in ranges_to_remove:
        ws_tpl.unmerge_cells(r)

    ws_tpl.delete_rows(7, 5)

    n = len(participants)
    if n > 0:
        ws_tpl.insert_rows(7, n)

    thin_side = Side(style='thin', color="000000")
    border_all = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    font_normal = Font(name='Times New Roman', size=12)
    fill_white = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center', wrap_text=True, indent=1)

    for i, p in enumerate(participants):
        row_idx = 7 + i
        ws_tpl.row_dimensions[row_idx].height = row7_height

        for col_idx in range(1, 11):
            cell = ws_tpl.cell(row=row_idx, column=col_idx)
            style = row7_styles[col_idx - 1]
            if style["font"]: cell.font = copy.copy(style["font"])
            if style["fill"]: cell.fill = copy.copy(style["fill"])
            if style["border"]: cell.border = copy.copy(style["border"])
            if style["alignment"]: cell.alignment = copy.copy(style["alignment"])
            cell.number_format = style["number_format"]

            cell.border = border_all
            cell.font = font_normal
            cell.fill = fill_white
            if col_idx in [1, 2, 5, 6, 7, 8]:
                cell.alignment = align_center
            else:
                cell.alignment = align_left

        ws_tpl.cell(row=row_idx, column=1).value = i + 1
        ws_tpl.cell(row=row_idx, column=2).value = p["danh_xung"]
        ws_tpl.cell(row=row_idx, column=3).value = p["ho_ten"]
        ws_tpl.cell(row=row_idx, column=4).value = p["cong_ty"]
        ws_tpl.cell(row=row_idx, column=5).value = p["cccd"]

        bday = p["ngay_sinh"]
        if isinstance(bday, datetime):
            bday_str = f"{bday.day:02d}/{bday.month:02d}/{bday.year}"
        elif bday:
            bday_str = str(bday)
        else:
            bday_str = ""
        ws_tpl.cell(row=row_idx, column=6).value = bday_str
        ws_tpl.cell(row=row_idx, column=7).value = p["gioi_tinh"]
        ws_tpl.cell(row=row_idx, column=8).value = p["sdt"]
        ws_tpl.cell(row=row_idx, column=9).value = p["email"]
        ws_tpl.cell(row=row_idx, column=10).value = ""

    # Re-append Lưu ý section
    luuy_row = 7 + n
    notes_row = luuy_row + 1

    ws_tpl.row_dimensions[luuy_row].height = luuy_header_height
    ws_tpl.cell(row=luuy_row, column=1).value = "Lưu ý"
    ws_tpl.cell(row=luuy_row, column=1).font = luuy_header_font
    ws_tpl.cell(row=luuy_row, column=1).alignment = luuy_header_align
    ws_tpl.merge_cells(start_row=luuy_row, start_column=1, end_row=luuy_row, end_column=10)

    ws_tpl.row_dimensions[notes_row].height = luuy_notes_height
    ws_tpl.cell(row=notes_row, column=1).value = notes_text
    ws_tpl.cell(row=notes_row, column=1).font = luuy_notes_font
    ws_tpl.cell(row=notes_row, column=1).alignment = luuy_notes_align
    ws_tpl.merge_cells(start_row=notes_row, start_column=1, end_row=notes_row, end_column=10)

    xlsx_buf = io.BytesIO()
    wb_tpl.save(xlsx_buf)
    xlsx_buf.seek(0)
    return xlsx_buf


@app.post("/api/training-dossier/export")
async def td_export_course(
    excel: UploadFile = File(...),
    template_xlsx: UploadFile = File(...),
    template_docx: UploadFile = File(...),
    course_json: str = Form(...),
):
    """Xuất hồ sơ đào tạo: tạo xlsx + docx theo từng đơn vị rồi nén thành ZIP trả về."""
    logs = []
    def log(type_, text):
        logs.append({"type": type_, "text": text})

    try:
        course = json.loads(course_json)
        excel_bytes = await excel.read()
        tpl_xlsx_bytes = await template_xlsx.read()
        tpl_docx_bytes = await template_docx.read()

        log("info", "📚 Đang đọc dữ liệu từ Excel...")
        _, _, rows, header, col_map = _td_parse_excel(excel_bytes)

        IDX_TEN_LOP = col_map.get("Khóa học", col_map.get("Tên lớp học", -1))
        IDX_NGAY_BD = col_map.get("Thời gian bắt đầu", col_map.get("Ngày bắt đầu", -1))
        IDX_NGAY_KT = col_map.get("Thời gian kết thúc", col_map.get("Ngày kết thúc", -1))
        IDX_HO_TEN = col_map.get("Họ và tên", -1)
        IDX_DANH_XUNG = col_map.get("Danh xưng", -1)
        IDX_NGAY_SINH = col_map.get("Ngày sinh", -1)
        IDX_CCCD = col_map.get("CCCD", -1)
        IDX_GIOI_TINH = col_map.get("Giới tính", -1)
        IDX_SDT = col_map.get("Số điện thoại", -1)
        IDX_EMAIL = col_map.get("Email", -1)
        IDX_DON_VI = col_map.get("Đơn vị", col_map.get("Đơn Vị", -1))
        IDX_GIANG_VIEN = col_map.get("Giảng viên", -1)

        # Filter participants for the selected course
        participants = []
        target_ten_clean = _td_clean_for_match(course["ten_lop"])
        for row in rows[1:]:
            if len(row) <= max(IDX_TEN_LOP, IDX_NGAY_BD): continue

            r_ten_clean = _td_clean_for_match(row[IDX_TEN_LOP])
            rbd = _td_parse_date(row[IDX_NGAY_BD])
            rbd_str = rbd.strftime("%Y-%m-%d") if rbd else ""

            if r_ten_clean == target_ten_clean and rbd_str == course["ngay_bd"]:
                def _safe_get(idx):
                    if idx == -1 or idx >= len(row) or row[idx] is None: return ""
                    return _td_normalize_text(row[idx])

                ho_ten = _safe_get(IDX_HO_TEN)
                if not ho_ten: continue

                participants.append({
                    "ho_ten": ho_ten,
                    "bo_phan": _safe_get(IDX_DON_VI),
                    "danh_xung": _safe_get(IDX_DANH_XUNG),
                    "ngay_sinh": _td_parse_date(row[IDX_NGAY_SINH]) if IDX_NGAY_SINH != -1 and IDX_NGAY_SINH < len(row) else None,
                    "cong_ty": "Công ty Cổ phần Phòng thử nghiệm Phúc Gia",
                    "cccd": _safe_get(IDX_CCCD),
                    "gioi_tinh": _safe_get(IDX_GIOI_TINH),
                    "sdt": _safe_get(IDX_SDT),
                    "email": _safe_get(IDX_EMAIL),
                })

        if not participants:
            return JSONResponse({"logs": logs, "error": "Không tìm thấy học viên nào cho khóa học này!"}, status_code=400)

        log("info", f"✓ Tìm thấy {len(participants)} học viên.")

        dt_bd = datetime.strptime(course["ngay_bd"], "%Y-%m-%d")
        dt_kt = datetime.strptime(course["ngay_kt"], "%Y-%m-%d") if course.get("ngay_kt") else dt_bd
        date_prefix = dt_bd.strftime("%Y.%m.%d")
        safe_name = _td_make_filename_safe(course["ten_lop"])

        # ── NHÓM PARTICIPANTS THEO ĐƠN VỊ ──
        from collections import OrderedDict
        unit_groups = OrderedDict()
        for p in participants:
            unit_key = _td_classify_unit(p["bo_phan"])
            if unit_key not in unit_groups:
                unit_groups[unit_key] = []
            unit_groups[unit_key].append(p)

        log("info", f"📋 Phân loại theo đơn vị: {', '.join(f'{k} ({len(v)} HV)' for k, v in unit_groups.items())}")

        # ── TẠO FILE CHO TỪNG ĐƠN VỊ + FILE ĐẦY ĐỦ ──
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:

            if len(unit_groups) == 1:
                # Chỉ có 1 đơn vị -> Không cần tạo thư mục con, xuất thẳng file vào ZIP
                unit_key = list(unit_groups.keys())[0]
                unit_safe = _td_make_filename_safe(unit_key)
                
                log("info", f"📝 Chỉ có 1 đơn vị ({unit_key}), xuất thẳng file...")
                docx_filename = f"{date_prefix}-PGE-TT11-BM03-{safe_name}-{unit_safe}.docx"
                docx_buf = _td_generate_docx(tpl_docx_bytes, course, participants, dt_bd, dt_kt)
                zf.writestr(docx_filename, docx_buf.getvalue())
                log("info", f"  ✓ {docx_filename}")

                xlsx_filename = f"{date_prefix}-{safe_name}-{unit_safe}.xlsx"
                xlsx_buf = _td_generate_xlsx(tpl_xlsx_bytes, course, participants)
                zf.writestr(xlsx_filename, xlsx_buf.getvalue())
                log("info", f"  ✓ {xlsx_filename}")

            else:
                # Có nhiều đơn vị -> Tạo thư mục cho từng đơn vị và thư mục Chung_Chi
                # 1) File cho từng đơn vị
                for unit_key, unit_participants in unit_groups.items():
                    unit_safe = _td_make_filename_safe(unit_key)
                    unit_folder = f"{unit_safe}/"

                    # DOCX cho đơn vị
                    log("info", f"📝 Đang tạo BM03 cho {unit_key} ({len(unit_participants)} HV)...")
                    unit_docx_filename = f"{date_prefix}-PGE-TT11-BM03-{safe_name}-{unit_safe}.docx"
                    unit_docx_buf = _td_generate_docx(tpl_docx_bytes, course, unit_participants, dt_bd, dt_kt)
                    zf.writestr(f"{unit_folder}{unit_docx_filename}", unit_docx_buf.getvalue())
                    log("info", f"  ✓ {unit_folder}{unit_docx_filename}")

                    # XLSX cho đơn vị
                    log("info", f"📊 Đang tạo danh sách cho {unit_key}...")
                    unit_xlsx_filename = f"{date_prefix}-{safe_name}-{unit_safe}.xlsx"
                    unit_xlsx_buf = _td_generate_xlsx(tpl_xlsx_bytes, course, unit_participants)
                    zf.writestr(f"{unit_folder}{unit_xlsx_filename}", unit_xlsx_buf.getvalue())
                    log("info", f"  ✓ {unit_folder}{unit_xlsx_filename}")

                # 2) File đầy đủ cho chứng chỉ
                full_folder = "Chung_Chi/"
                full_docx_filename = f"{date_prefix}-PGE-TT11-BM03-{safe_name}.docx"
                full_xlsx_filename = f"{date_prefix}-{safe_name}.xlsx"

                log("info", f"📝 Đang tạo file đầy đủ ({len(participants)} HV) cho chứng chỉ...")
                full_docx_buf = _td_generate_docx(tpl_docx_bytes, course, participants, dt_bd, dt_kt)
                zf.writestr(f"{full_folder}{full_docx_filename}", full_docx_buf.getvalue())
                log("info", f"  ✓ {full_folder}{full_docx_filename}")

                full_xlsx_buf = _td_generate_xlsx(tpl_xlsx_bytes, course, participants)
                zf.writestr(f"{full_folder}{full_xlsx_filename}", full_xlsx_buf.getvalue())
                log("info", f"  ✓ {full_folder}{full_xlsx_filename}")

        zip_buf.seek(0)
        zip_name = f"{date_prefix}-Ho_So_Dao_Tao-{safe_name}.zip"
        log("success", f"✅ Xong! Đã tạo hồ sơ đào tạo cho {len(participants)} học viên ({len(unit_groups)} đơn vị).")

        return Response(
            content=zip_buf.getvalue(),
            media_type="application/zip",
            headers={
                "Content-Disposition": f'attachment; filename="{zip_name}"',
                "x-file-count": str(len(participants)),
                "x-logs": str(len(logs)),
            }
        )

    except json.JSONDecodeError:
        return JSONResponse({"logs": logs, "error": "Dữ liệu khóa học không hợp lệ (JSON)."}, status_code=400)
    except Exception as e:
        import traceback
        log("error", f"❌ Lỗi: {str(e)}")
        log("error", traceback.format_exc())
        return JSONResponse({"logs": logs, "error": str(e)}, status_code=500)


if __name__ == "__main__":
    import uvicorn
    # Render provides the PORT environment variable
    port = int(os.environ.get("PORT", 8000))
    print("=" * 50)
    print(f"PGE Backend Server starting on port {port}...")
    print("=" * 50)
    uvicorn.run("backend:app", host="0.0.0.0", port=port, reload=False)

